#!/usr/bin/env python3
"""Generate Infinity Clinic User Guide DOCX."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = "Infinity-Clinic-User-Guide.docx"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)
    note = doc.add_paragraph(
        "Tip: In Microsoft Word, right-click the table of contents and choose Update Field."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def add_table(doc, headers, rows, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], header_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        if level:
            p.paragraph_format.left_indent = Inches(0.25 * level)


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Infinity Clinic\nManagement System")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("User Guide")
    sub_run.font.size = Pt(20)
    sub_run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Version 1.0  |  {datetime.date.today().strftime('%B %d, %Y')}\n"
        "Covers: Public website, Receptionist, Doctor, Pharmacist, and Admin portals"
    ).font.size = Pt(11)

    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Infinity Clinic is a clinic management platform that connects the public patient website, "
        "reception desk, doctor consultation room, pharmacy desk, and admin control panel. "
        "This guide explains how to use each part of the system for your role."
    )
    doc.add_heading("1.1 Who This Guide Is For", level=2)
    add_bullets(doc, [
        "Patients and visitors — booking appointments on the public website (no login required)",
        "Receptionists — managing appointments, check-in, OPD tokens, walk-ins, and payments",
        "Doctors — running the OPD queue, consultations, and prescriptions",
        "Pharmacists — dispensing medicines from the live prescription queue",
        "Administrators — staff management, website content, clinic settings, and oversight",
    ])
    doc.add_heading("1.2 System Overview", level=2)
    doc.add_paragraph(
        "A typical patient visit flows through these steps:"
    )
    add_numbered(doc, [
        "Patient books online or walks in at reception",
        "Receptionist confirms the appointment (if needed) and checks the patient in",
        "An OPD token is issued and the patient appears in the doctor's live queue",
        "Doctor calls the patient, records consultation notes, and optionally writes a prescription",
        "If medicines are prescribed, the prescription appears in the pharmacy queue",
        "Pharmacist dispenses medicines",
        "Receptionist records offline payment and prints a receipt",
    ])
    doc.add_paragraph(
        "Live updates for the OPD queue and pharmacy queue happen in real time — "
        "no manual page refresh is needed when Socket.IO is connected."
    )

    doc.add_page_break()

    # Staff login
    doc.add_heading("2. Staff Login (All Roles)", level=1)
    doc.add_paragraph("All staff members use the same login page.")
    add_table(doc, ["Item", "Details"], [
        ["Login URL", "/portal/login (or click Staff on the public website)"],
        ["Page title", "Staff Portal"],
        ["Fields", "Email and Password"],
        ["Sign in button", "Sign In"],
        ["After login", "You are redirected to your role's dashboard"],
        ["Logout", "Log out — in the sidebar (desktop) or mobile menu"],
    ])
    doc.add_heading("2.1 Demo Login Credentials", level=2)
    doc.add_paragraph(
        "The following accounts are available in the demo/seed environment. "
        "Change all passwords before going live."
    )
    add_table(doc, ["Role", "Email", "Password"], [
        ["Admin", "admin@infinityclinic.com", "Admin@123"],
        ["Doctor (Cardiology)", "doctor@infinityclinic.com", "Doctor@123"],
        ["Receptionist", "receptionist@infinityclinic.com", "Reception@123"],
        ["Pharmacist", "pharmacy@infinityclinic.com", "Pharmacy@123"],
    ])
    doc.add_heading("2.2 Additional Doctor Accounts (Demo)", level=2)
    add_table(doc, ["Doctor", "Email", "Password", "Specialization"], [
        ["Dr. Prasann Moon", "moon@infinityclinics.com", "Doctor@123", "ENT"],
        ["Dr. Gunjan Kolhe", "kolhe@infinityclinics.com", "Doctor@123", "Orthopaedics"],
        ["Dr. Pranit Khandait", "khandait@infinityclinics.com", "Doctor@123", "Neurology"],
        ["Dr. Shweta Lodhi", "lodhi@infinityclinics.com", "Doctor@123", "Gynaecology"],
    ])

    doc.add_page_break()

    # Public / Patient
    doc.add_heading("3. Public Website & Online Booking (Patients)", level=1)
    doc.add_paragraph(
        "Patients do not need an account. The public website is for browsing clinic information "
        "and booking appointments online."
    )
    doc.add_heading("3.1 Website Pages", level=2)
    add_table(doc, ["Page", "URL", "Purpose"], [
        ["Home", "/", "Clinic overview, departments, services, testimonials, location"],
        ["About", "/about", "Clinic story, values, first-visit guidance"],
        ["Specialists", "/doctors", "Doctor profiles, fees, timings, book links"],
        ["Services", "/services", "Department services and diagnostics"],
        ["Visit Us", "/contact", "Phone, WhatsApp, address, timings, FAQ"],
        ["Patient Stories", "/testimonials", "Patient testimonials"],
        ["Book Online", "/book", "Appointment booking wizard"],
    ])
    doc.add_heading("3.2 How to Book an Appointment Online", level=2)
    add_numbered(doc, [
        "Go to Book Online (/book) or click Book with [Doctor Name] from the Specialists page.",
        "Step 1 — Doctor: Select a doctor from the list. Each card shows name, specialization, and consultation fee.",
        "Step 2 — Date: Pick an available date on the calendar. Past dates cannot be selected.",
        "Step 3 — Time: Choose an available time slot. If no slots appear, try another date or call the clinic.",
        "Step 4 — Details: Enter Full Name, Phone, and optional Email. Review the summary.",
        "Click Confirm Booking.",
        "Step 5 — Done: You will see Appointment Confirmed with doctor, date, and time details.",
    ])
    doc.add_paragraph(
        "Online bookings are created with status Pending. A receptionist must confirm the appointment "
        "before the visit day (or on arrival)."
    )
    doc.add_heading("3.3 What Patients Should Know", level=2)
    add_bullets(doc, [
        "There is no patient login portal — you cannot view records or prescriptions online.",
        "Bring your phone number used during booking for faster check-in at reception.",
        "For walk-in visits, go directly to the reception desk — no online booking needed.",
        "Payment is collected at the clinic (cash, card, or UPI) — there is no online payment gateway.",
    ])

    doc.add_page_break()

    # Receptionist
    doc.add_heading("4. Receptionist Portal", level=1)
    doc.add_paragraph(
        "Receptionists manage the front desk: appointments, patient records, check-in, "
        "OPD tokens, walk-ins, and payment recording."
    )
    doc.add_heading("4.1 Navigation", level=2)
    add_table(doc, ["Menu", "URL", "Purpose"], [
        ["Dashboard", "/portal/receptionist", "Today's overview, appointments, live queue"],
        ["Appointments", "/portal/receptionist/appointments", "Full appointment management"],
        ["Patients", "/portal/receptionist/patients", "Search, add, edit patients and visit history"],
        ["Walk-in", "/portal/receptionist/walk-in", "Register patients without prior booking"],
    ])
    doc.add_heading("4.2 Dashboard", level=2)
    doc.add_paragraph("The Reception Desk dashboard shows:")
    add_bullets(doc, [
        "Metrics: Today's Total, Awaiting Check-in, In Clinic, Completed",
        "Today's Appointments table with search and status filter",
        "Live Queue — select a doctor to see their real-time OPD queue",
        "Quick actions: + Walk-in, All Appointments",
    ])
    doc.add_heading("4.3 Appointment Management", level=2)
    doc.add_paragraph("Available actions depend on appointment status:")
    add_table(doc, ["Status", "Meaning", "Available Actions"], [
        ["Pending", "Booked online, not yet confirmed", "View, Confirm, Reschedule, Cancel, No Show"],
        ["Confirmed", "Expected visit, not arrived", "View, Check In, Reschedule, Cancel, No Show"],
        ["Checked In", "Arrived, OPD token issued", "View, Payment"],
        ["In Consultation", "Currently with doctor", "View, Payment"],
        ["Completed", "Visit finished", "View, Payment"],
        ["Cancelled", "Appointment cancelled", "View only"],
        ["No Show", "Patient did not arrive", "View only"],
    ])
    doc.add_heading("4.4 Check-In Workflow", level=2)
    add_numbered(doc, [
        "Find the patient's appointment on the Dashboard or Appointments page.",
        "Click Check In.",
        "An OPD token number is automatically issued for the assigned doctor.",
        "The patient appears in the doctor's live queue with status Waiting.",
        "Direct the patient to the waiting area.",
    ])
    doc.add_heading("4.5 Walk-In Workflow", level=2)
    add_numbered(doc, [
        "Click + Walk-in (or go to the Walk-in page).",
        "Select the doctor.",
        "Enter patient phone and name (creates or matches existing patient).",
        "Optionally add notes.",
        "Click Register Walk-in.",
        "Return to the Dashboard and Check In the walk-in appointment to issue a token.",
    ])
    doc.add_heading("4.6 Patient Management", level=2)
    add_bullets(doc, [
        "Search patients by name or phone (minimum 2 characters).",
        "Add Patient — create a new patient record (name, phone, email, DOB, gender, address).",
        "Edit — update patient details.",
        "View — open patient profile with visit history.",
        "Visit History — click View on any past visit to see consultation details in a modal.",
    ])
    doc.add_heading("4.7 Recording Payment", level=2)
    add_numbered(doc, [
        "After the visit is completed (or while in consultation), click Payment on the appointment row.",
        "Enter the amount in rupees (₹).",
        "Select payment method: Cash, Card, or UPI.",
        "Click Record & Print Receipt.",
        "Use Print to print the receipt, then Close.",
    ])
    doc.add_paragraph(
        "Note: This records offline payments only. There is no online payment gateway integrated."
    )
    doc.add_heading("4.8 Reschedule & Cancel", level=2)
    add_bullets(doc, [
        "Reschedule — opens a modal to pick a new date and time, then click Reschedule.",
        "Cancel — cancels the appointment (status becomes Cancelled).",
        "No Show — marks the patient as not having arrived.",
    ])
    doc.add_heading("4.9 Live Queue", level=2)
    doc.add_paragraph(
        "The Live Queue panel shows real-time OPD tokens for the selected doctor. "
        "Token statuses:"
    )
    add_table(doc, ["Token Status", "Meaning"], [
        ["Waiting", "Patient is in queue, not yet called"],
        ["Called", "Doctor has called this token"],
        ["In Consultation", "Patient is with the doctor"],
        ["Completed", "Consultation finished"],
        ["Skipped", "Doctor skipped this token (can be recalled)"],
    ])

    doc.add_page_break()

    # Doctor
    doc.add_heading("5. Doctor Portal", level=1)
    doc.add_paragraph(
        "Doctors use the portal to manage their OPD queue, conduct consultations, "
        "write prescriptions, and review patient history."
    )
    doc.add_heading("5.1 Navigation", level=2)
    add_table(doc, ["Menu", "URL", "Purpose"], [
        ["Dashboard", "/portal/doctor", "Consultation room — live queue and current patient"],
        ["Appointments", "/portal/doctor/appointments", "Today's scheduled appointments"],
        ["Patients", "/portal/doctor/patients", "Search and view patient profiles"],
        ["History", "/portal/doctor/history", "Past consultations (7/30/90 days)"],
    ])
    doc.add_heading("5.2 Consultation Room (Dashboard)", level=2)
    doc.add_paragraph("The main dashboard is your Consultation Room:")
    add_bullets(doc, [
        "OPD Queue — real-time list of waiting patients with token numbers",
        "Search and filter queue by status",
        "Completed and Waiting counts at the bottom",
        "Current Patient panel — appears when you start a consultation",
    ])
    doc.add_heading("5.3 Consultation Workflow", level=2)
    add_numbered(doc, [
        "A checked-in patient appears in your queue with status Waiting.",
        "Click Call to announce the token (status → Called), or click Start to begin directly.",
        "The Current Patient panel opens with consultation fields.",
        "Fill in Chief Complaint, Diagnosis, and Clinical Notes.",
        "Optionally add a prescription (see section 5.4).",
        "Click Complete Visit & Send Rx (if medicines added) or Complete Visit (No Rx).",
        "The patient is removed from the active queue and the next patient can be called.",
    ])
    doc.add_heading("5.4 Writing a Prescription", level=2)
    doc.add_paragraph("Prescriptions are optional. You can complete a visit without prescribing medicines.")
    add_bullets(doc, [
        "Saved medicines — click a chip to quickly add a commonly used medicine.",
        "For each medicine: Medicine name, Dose, Times per day, Duration, When to take, Extra instructions.",
        "Click + Add medicine to add more rows.",
        "Click Remove to delete a medicine row.",
        "General advice — free-text advice shown on the prescription.",
        "Print Rx — print the prescription (available after saving medicines).",
    ])
    doc.add_paragraph("When you complete the visit:")
    add_bullets(doc, [
        "With medicines: prescription is sent to the pharmacy queue automatically.",
        "Without medicines: visit completes normally; pharmacy is not notified.",
        "Advice-only visits (rest, follow-up, lifestyle) work without adding medicine rows.",
    ])
    doc.add_heading("5.5 Queue Actions", level=2)
    add_table(doc, ["Action", "When Available", "What It Does"], [
        ["Call", "Token status: Waiting", "Announces the token (status → Called)"],
        ["Start", "Waiting or Called", "Opens consultation form (status → In Consultation)"],
        ["Skip", "Waiting", "Skips this token (patient can be recalled later)"],
        ["View", "Any", "Opens patient profile page"],
    ])
    doc.add_heading("5.6 Patient Profiles & History", level=2)
    add_bullets(doc, [
        "Patients page — search by name or phone, click View Profile.",
        "Patient detail — shows contact info and Your visits (past consultations with this doctor).",
        "Each visit card shows complaint, diagnosis, prescription summary, and Print Rx.",
        "History page — table of past consultations with View (modal) and Patient links.",
        "Use Back to return to the previous page.",
    ])

    doc.add_page_break()

    # Pharmacist
    doc.add_heading("6. Pharmacist Portal", level=1)
    doc.add_paragraph(
        "Pharmacists manage the live prescription queue and record when medicines have been dispensed."
    )
    doc.add_heading("6.1 Navigation", level=2)
    add_table(doc, ["Menu", "URL", "Purpose"], [
        ["Queue", "/portal/pharmacy", "Live prescription queue"],
        ["History", "/portal/pharmacy/history", "Dispensed prescriptions history"],
    ])
    doc.add_heading("6.2 Pharmacy Desk (Queue)", level=2)
    add_bullets(doc, [
        "Metrics: Waiting, In Progress, Total in Queue",
        "Live Prescription Queue — updates in real time when doctors complete visits",
        "Filter by All, Waiting, or In progress",
        "Click a row or View to open prescription details",
    ])
    doc.add_heading("6.3 Dispensing Workflow", level=2)
    add_numbered(doc, [
        "When a doctor completes a visit with medicines, the prescription appears in the queue.",
        "Click the prescription to view patient name, doctor, token number, medicines, and advice.",
        "Opening a waiting prescription automatically marks it In Progress.",
        "Prepare and hand over the medicines to the patient.",
        "Click Mark as Dispensed when done.",
        "The prescription moves to Dispensed History.",
    ])
    doc.add_paragraph(
        "If a doctor completes a visit without prescribing medicines, nothing appears in the pharmacy queue."
    )
    doc.add_heading("6.4 Dispensed History", level=2)
    add_bullets(doc, [
        "Filter by Today, Last 7 days, or Last 30 days.",
        "Table shows dispensed time, patient, doctor, medicine count, and token.",
        "Click View to open the full prescription in a modal.",
    ])
    doc.add_heading("6.5 Prescription Statuses", level=2)
    add_table(doc, ["Status", "Meaning"], [
        ["Waiting (pending)", "New prescription, not yet opened"],
        ["In Progress (dispensing)", "Pharmacist is preparing medicines"],
        ["Dispensed", "Medicines given to patient"],
    ])

    doc.add_page_break()

    # Admin
    doc.add_heading("7. Admin Portal", level=1)
    doc.add_paragraph(
        "Administrators have full access to clinic oversight, staff management, "
        "website content, settings, and permissions."
    )
    doc.add_heading("7.1 Navigation", level=2)
    add_table(doc, ["Menu", "URL", "Purpose"], [
        ["Dashboard", "/portal/admin", "Clinic metrics and quick links"],
        ["Doctors", "/portal/admin/doctors", "Add doctors, schedules, deactivate"],
        ["Receptionists", "/portal/admin/receptionists", "Add and manage receptionists"],
        ["Appointments", "/portal/admin/appointments", "View all appointments (read-only)"],
        ["Website CMS", "/portal/admin/cms", "Edit public website content"],
        ["Settings", "/portal/admin/settings", "Clinic name, contact, booking settings"],
        ["Permissions", "/portal/admin/permissions", "Role permission matrix"],
    ])
    doc.add_paragraph(
        "Admins also have quick links to open Receptionist View, Doctor View, and Pharmacy View "
        "to test or assist with workflows."
    )
    doc.add_heading("7.2 Dashboard", level=2)
    add_bullets(doc, [
        "Total Patients, Active Doctors, Today's Appointments, Completed Today",
        "Revenue (last 30 days) from recorded payments",
        "Appointments by Status chart (last 30 days)",
        "Quick Links to all major admin and portal pages",
    ])
    doc.add_heading("7.3 Managing Doctors", level=2)
    add_numbered(doc, [
        "Go to Doctors and click Add Doctor.",
        "Fill in: Email, Password, Full Name, Specialization, Qualification, Consultation Fee, Bio.",
        "Click Create Doctor.",
        "Select a doctor to view their profile and weekly schedule.",
        "Add schedule slots: Day (Sun–Sat), Start time, End time → Add Slot.",
        "Deactivate to disable a doctor (they will not appear for booking).",
    ])
    doc.add_heading("7.4 Managing Receptionists", level=2)
    add_numbered(doc, [
        "Go to Receptionists and click Add Receptionist.",
        "Enter Email, Password, and Full Name.",
        "Click Create.",
        "Use Deactivate to disable an account.",
    ])
    doc.add_paragraph(
        "Note: Pharmacist accounts are created via system setup/seed only. "
        "There is no admin UI to add pharmacists in the current version."
    )
    doc.add_heading("7.5 Appointments (Oversight)", level=2)
    doc.add_paragraph(
        "The admin Appointments page is read-only. Use filters for date, doctor, status, and search. "
        "Click View to see appointment details. For check-in, payments, and queue management, "
        "use the Receptionist View."
    )
    doc.add_heading("7.6 Website CMS", level=2)
    doc.add_paragraph("Edit all public website content without touching code. Tabs:")
    add_table(doc, ["Tab", "What You Can Edit"], [
        ["Pages", "Hero, About, Doctors page, Services page, Contact page, Testimonials page, Book page"],
        ["Site", "Home, Contact, CTA banner, Footer"],
        ["Lists", "Why choose us cards, Visit steps, FAQ, Diagnostics (JSON)"],
        ["Services", "Add, view, delete service entries"],
        ["Testimonials", "Add, view, delete patient testimonials"],
    ])
    doc.add_paragraph(
        "Click Save on each section. Changes appear on the public website immediately."
    )
    doc.add_heading("7.7 Clinic Settings", level=2)
    add_bullets(doc, [
        "Clinic Identity: Clinic Name, Phone, Email, Address",
        "Booking: Default Slot Duration (minutes)",
        "Click Save Settings — updates the database and public contact section",
    ])
    doc.add_heading("7.8 Role Permissions", level=2)
    doc.add_paragraph(
        "Configure what receptionists and doctors can do. Admin always has full access."
    )
    add_bullets(doc, [
        "Permission groups: Reception Desk, Patients, OPD Queue, Payments, Consultations, "
        "Prescriptions, Pharmacy, Staff Management, Website & CMS, Clinic Settings, Portal Access",
        "Reset defaults — restore factory permission settings",
        "Save permissions — staff may need to log out and back in to see changes",
    ])

    doc.add_page_break()

    # Reference
    doc.add_heading("8. Quick Reference", level=1)
    doc.add_heading("8.1 Appointment Statuses", level=2)
    add_table(doc, ["Status", "Description"], [
        ["Pending", "Booked online, awaiting confirmation"],
        ["Confirmed", "Confirmed, patient expected"],
        ["Checked In", "Patient arrived, OPD token issued"],
        ["In Consultation", "Patient is with the doctor"],
        ["Completed", "Visit finished"],
        ["Cancelled", "Appointment cancelled"],
        ["No Show", "Patient did not arrive"],
    ])
    doc.add_heading("8.2 OPD Token Statuses", level=2)
    add_table(doc, ["Status", "Description"], [
        ["Waiting", "In queue, not yet called"],
        ["Called", "Doctor announced this token"],
        ["In Consultation", "With the doctor"],
        ["Completed", "Consultation done"],
        ["Skipped", "Passed over by doctor"],
    ])
    doc.add_heading("8.3 Payment Methods", level=2)
    doc.add_paragraph("Cash, Card, UPI — all recorded offline at reception. No online payment gateway.")
    doc.add_heading("8.4 How Appointments Are Booked", level=2)
    add_table(doc, ["Source", "Description"], [
        ["website", "Patient booked online at /book"],
        ["walk_in", "Registered at reception walk-in desk"],
        ["phone", "Booked by phone (recorded by receptionist)"],
    ])
    doc.add_heading("8.5 Real-Time Updates", level=2)
    add_bullets(doc, [
        "OPD queue — updates live on Receptionist and Doctor dashboards",
        "Pharmacy queue — updates live on Pharmacy dashboard",
        "Requires Redis and Socket.IO to be running on the server",
        "On reconnect, the system automatically refreshes the latest queue snapshot",
    ])

    doc.add_page_break()

    # Troubleshooting
    doc.add_heading("9. Troubleshooting & Tips", level=1)
    doc.add_heading("9.1 Common Issues", level=2)
    add_table(doc, ["Problem", "Solution"], [
        ["Queue not updating live", "Check that Redis is running. Refresh the page. Log out and back in."],
        ["Cannot check in patient", "Appointment must be Pending or Confirmed. Cancelled/No Show cannot be checked in."],
        ["No time slots on booking page", "Doctor may have no schedule for that day, or all slots are taken. Try another date."],
        ["Prescription not in pharmacy queue", "Doctor must add medicines and click Complete Visit & Send Rx. Advice-only visits do not create pharmacy entries."],
        ["Permission denied error", "Admin may have restricted your role. Contact admin or log out and back in after permission changes."],
        ["Forgot password", "Contact your clinic administrator to reset your password."],
    ])
    doc.add_heading("9.2 Best Practices", level=2)
    add_bullets(doc, [
        "Reception: Confirm online bookings the same day or on arrival before check-in.",
        "Reception: Always check in patients before directing them to the waiting area.",
        "Doctor: Use Call before Start so reception and display boards can track called tokens.",
        "Doctor: Complete visits promptly so the pharmacy queue stays accurate.",
        "Pharmacy: Mark prescriptions as dispensed only after handing medicines to the patient.",
        "Admin: Keep doctor schedules up to date so online booking shows correct slots.",
        "Admin: Review CMS content regularly and update clinic contact details in Settings.",
    ])

    doc.add_page_break()

    # Appendix
    doc.add_heading("Appendix A: End-to-End Visit Example", level=1)
    add_numbered(doc, [
        "Patient visits infinityclinic.com and books Dr. Maske for tomorrow at 10:00 AM.",
        "Booking status: Pending.",
        "Next day — receptionist opens Dashboard, finds the appointment, clicks Confirm.",
        "Patient arrives. Receptionist clicks Check In. Token #7 is issued.",
        "Dr. Maske sees Token #7 in his OPD queue. He clicks Call, then Start.",
        "He records: Chief complaint 'Chest pain', Diagnosis 'Mild gastritis', adds Pantoprazole 40mg for 14 days.",
        "He clicks Complete Visit & Send Rx.",
        "Pharmacy desk sees the new prescription. Pharmacist dispenses medicines and clicks Mark as Dispensed.",
        "Receptionist clicks Payment on the appointment, records ₹500 Cash, prints receipt.",
        "Visit complete.",
    ])

    doc.add_heading("Appendix B: Glossary", level=1)
    add_table(doc, ["Term", "Definition"], [
        ["OPD", "Outpatient Department — same-day clinic visits without admission"],
        ["Token", "Queue number issued at check-in for the doctor's waiting list"],
        ["Check-in", "Reception action that marks patient arrival and issues a token"],
        ["Walk-in", "Patient who arrives without a prior online booking"],
        ["CMS", "Content Management System — admin tool to edit website text and images"],
        ["Rx / Prescription", "Doctor's medicine order sent to pharmacy after consultation"],
        ["Socket.IO", "Technology used for live queue updates without page refresh"],
    ])

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
